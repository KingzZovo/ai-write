"""ExportService -- render a Project (+ volumes + chapters) as EPUB / PDF / DOCX.

Public entry points return ``(bytes, filename, mime)`` tuples so the API layer
can stream them back as a download. Each builder loads the project with its
volumes and chapters ordered by ``volume_idx`` and ``chapter_idx``.

Kept deliberately dependency-light:
  - EPUB: ebooklib (already in pyproject)
  - PDF: reportlab (pure Python, added in chunk-13)
  - DOCX: python-docx (pure Python, added in chunk-13)
"""

from __future__ import annotations

import io
import logging
import re
import uuid
from typing import List, Tuple

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models.project import Chapter, Project, Volume

logger = logging.getLogger(__name__)

EXPORT_MIME = {
    "epub": "application/epub+zip",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


# ---------------------------------------------------------------------------
# Loading
# ---------------------------------------------------------------------------
async def _load_project_tree(db: AsyncSession, project_id: uuid.UUID) -> Project:
    stmt = (
        select(Project)
        .where(Project.id == project_id)
        .options(selectinload(Project.volumes).selectinload(Volume.chapters))
    )
    result = await db.execute(stmt)
    project = result.scalar_one_or_none()
    if project is None:
        raise LookupError(f"project not found: {project_id}")
    return project


def _ordered_chapters(project: Project) -> List[Tuple[Volume, Chapter]]:
    pairs: List[Tuple[Volume, Chapter]] = []
    for vol in sorted(project.volumes or [], key=lambda v: (v.volume_idx or 0)):
        for ch in sorted(vol.chapters or [], key=lambda c: (c.chapter_idx or 0)):
            pairs.append((vol, ch))
    return pairs


def _safe_filename(title: str, ext: str) -> str:
    base = re.sub(r"[^\w\u4e00-\u9fff\-\. ]+", "_", title or "untitled").strip(" ._") or "untitled"
    return f"{base[:80]}.{ext}"


# ---------------------------------------------------------------------------
# EPUB
# ---------------------------------------------------------------------------
def _build_epub_bytes(project: Project) -> bytes:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(f"ai-write:{project.id}")
    book.set_title(project.title or "Untitled")
    book.set_language("zh")
    book.add_author("ai-write")

    spine: list = ["nav"]
    toc: list = []
    for vol, ch in _ordered_chapters(project):
        ch_title = ch.title or f"第{ch.chapter_idx}章"
        body = (ch.content_text or "").replace("\r\n", "\n")
        paragraphs = "".join(f"<p>{_escape(p)}</p>" for p in body.split("\n") if p.strip())
        html = f"<h2>{_escape(ch_title)}</h2>{paragraphs}"
        item = epub.EpubHtml(
            title=ch_title,
            file_name=f"chap_{vol.volume_idx or 0:02d}_{ch.chapter_idx or 0:03d}.xhtml",
            lang="zh",
        )
        item.content = html
        book.add_item(item)
        spine.append(item)
        toc.append(item)

    book.toc = tuple(toc)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


def _escape(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _build_pdf_bytes(project: Project) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        PageBreak,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=project.title or "Untitled",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], fontSize=22, leading=28)
    ch_style = ParagraphStyle("ch", parent=styles["Heading2"], fontSize=16, leading=22)
    body_style = ParagraphStyle("body", parent=styles["BodyText"], fontSize=11, leading=18)

    story: list = []
    story.append(Paragraph(_escape(project.title or "Untitled"), title_style))
    story.append(Spacer(1, 1 * cm))

    for idx, (vol, ch) in enumerate(_ordered_chapters(project)):
        if idx > 0:
            story.append(PageBreak())
        story.append(Paragraph(_escape(ch.title or f"Chapter {ch.chapter_idx}"), ch_style))
        story.append(Spacer(1, 0.5 * cm))
        body = (ch.content_text or "").replace("\r\n", "\n")
        for para in body.split("\n"):
            if para.strip():
                story.append(Paragraph(_escape(para), body_style))
                story.append(Spacer(1, 0.2 * cm))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _build_docx_bytes(project: Project) -> bytes:
    from docx import Document

    doc = Document()
    doc.core_properties.title = project.title or "Untitled"
    doc.add_heading(project.title or "Untitled", level=0)

    current_volume_idx = None
    for vol, ch in _ordered_chapters(project):
        if vol.volume_idx != current_volume_idx:
            doc.add_heading(vol.title or f"Volume {vol.volume_idx}", level=1)
            current_volume_idx = vol.volume_idx
        doc.add_heading(ch.title or f"Chapter {ch.chapter_idx}", level=2)
        body = (ch.content_text or "").replace("\r\n", "\n")
        for para in body.split("\n"):
            if para.strip():
                doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
async def export_project(
    db: AsyncSession,
    *,
    project_id: uuid.UUID,
    fmt: str,
) -> Tuple[bytes, str, str]:
    """Render ``project_id`` as ``fmt`` ('epub' | 'pdf' | 'docx')."""
    fmt = fmt.lower().strip()
    if fmt not in EXPORT_MIME:
        raise ValueError(f"unsupported export format: {fmt!r}")

    project = await _load_project_tree(db, project_id)
    if fmt == "epub":
        data = _build_epub_bytes(project)
    elif fmt == "pdf":
        data = _build_pdf_bytes(project)
    else:  # docx
        data = _build_docx_bytes(project)

    filename = _safe_filename(project.title or str(project.id), fmt)
    return data, filename, EXPORT_MIME[fmt]
